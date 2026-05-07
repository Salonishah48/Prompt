
import io
import os
import re
from datetime import datetime, date
from collections import Counter

import numpy as np
import pandas as pd
import streamlit as st
from docx import Document

# ---------------------------
# Utility functions
# ---------------------------
STOPWORDS = set(
    ["the","a","an","and","or","but","if","to","for","with","on","in","of","by","is","are","be","as","at","from","this","that","these","those","it","into","over","under","between","within","about","your","you","we"]
)


def read_tabular(uploaded_file):
    """Read an uploaded Excel/CSV file into a pandas DataFrame.
    - For Excel, reads the first visible sheet by default.
    - Adds source_file column.
    """
    name = uploaded_file.name
    suffix = name.lower().split(".")[-1]

    try:
        if suffix in ["xlsx", "xls", "xlsm"]:
            # Use ExcelFile to safely handle different engines; openpyxl (xlsx), xlrd (xls)
            xls = pd.ExcelFile(uploaded_file)
            sheet_name = xls.sheet_names[0]
            # For xlsx/xlsm, prefer openpyxl; for xls, xlrd
            engine = "openpyxl" if suffix in ["xlsx", "xlsm"] else "xlrd"
            df = pd.read_excel(xls, sheet_name=sheet_name, engine=engine)
        elif suffix == "csv":
            df = pd.read_csv(uploaded_file)
        else:
            st.warning(f"Unsupported file type for {name}. Skipped.")
            return None
    except Exception as e:
        st.error(f"Failed to read {name}: {e}")
        return None

    df["source_file"] = name
    return df


def coerce_dates(series):
    """Try to convert a series to datetime; return converted series and a flag if successful."""
    try:
        converted = pd.to_datetime(series, errors="raise")
        return converted, True
    except Exception:
        try:
            converted = pd.to_datetime(series, errors="coerce")
            if converted.notna().sum() / max(1, len(series)) > 0.7:
                return converted, True
        except Exception:
            pass
    return series, False


def numeric_outlier_count(series):
    """Return count of outliers using IQR method for numeric series."""
    s = pd.to_numeric(series, errors="coerce")
    s = s.dropna()
    if s.empty:
        return 0
    q1 = s.quantile(0.25)
    q3 = s.quantile(0.75)
    iqr = q3 - q1
    lower = q1 - 1.5 * iqr
    upper = q3 + 1.5 * iqr
    return int(((s < lower) | (s > upper)).sum())


def profile_dataframe(df: pd.DataFrame):
    """Compute a compact profile of the dataframe for review points."""
    profile = {
        "rows": int(len(df)),
        "cols": int(df.shape[1]),
        "duplicates": int(df.duplicated().sum()),
        "columns": [],
        "files": sorted(df["source_file"].unique().tolist()) if "source_file" in df.columns else [],
    }

    for col in df.columns:
        if col == "source_file":
            continue
        s = df[col]
        nulls = int(s.isna().sum())
        null_pct = float(round(100 * nulls / max(1, len(s)), 2))

        # Detect date-like columns
        s_date, is_date = coerce_dates(s)

        # Numeric stats
        numeric_stats = None
        outliers = 0
        s_num = pd.to_numeric(s, errors="coerce")
        if s_num.notna().any():
            desc = s_num.describe()
            numeric_stats = {
                "count": int(desc.get("count", 0)),
                "mean": float(desc.get("mean", np.nan)) if not np.isnan(desc.get("mean", np.nan)) else None,
                "std": float(desc.get("std", np.nan)) if not np.isnan(desc.get("std", np.nan)) else None,
                "min": float(desc.get("min", np.nan)) if not np.isnan(desc.get("min", np.nan)) else None,
                "max": float(desc.get("max", np.nan)) if not np.isnan(desc.get("max", np.nan)) else None,
            }
            outliers = numeric_outlier_count(s)

        # Categorical samples
        vc = s.astype(str).str.slice(0, 80).value_counts(dropna=True)
        top_values = [f"{idx} ({int(cnt)})" for idx, cnt in vc.head(5).items()]

        col_info = {
            "name": col,
            "dtype": str(s.dtype),
            "nulls": nulls,
            "null_pct": null_pct,
            "unique": int(s.nunique(dropna=True)),
            "top_values": top_values,
            "is_date": is_date,
            "date_min": s_date.min().date().isoformat() if is_date and pd.notna(s_date.min()) else None,
            "date_max": s_date.max().date().isoformat() if is_date and pd.notna(s_date.max()) else None,
            "numeric_stats": numeric_stats,
            "outliers": outliers,
        }
        profile["columns"].append(col_info)

    # Potential keys (near-unique)
    n = len(df)
    profile["potential_keys"] = [c["name"] for c in profile["columns"] if c["unique"] >= max(1, int(0.98 * n))]

    return profile


def extract_keywords(prompt_text: str):
    if not prompt_text:
        return []
    words = re.findall(r"[A-Za-z0-9_]+", prompt_text.lower())
    return [w for w in words if len(w) > 2 and w not in STOPWORDS]


def prompt_aligned_highlights(df: pd.DataFrame, keywords):
    """Return a list of human-friendly highlights aligned to prompt keywords."""
    if not keywords:
        return []

    highlights = []
    cols = df.columns.tolist()
    # 1) Column alignment
    matched_cols = [c for c in cols if any(k in c.lower() for k in keywords)]
    if matched_cols:
        highlights.append(f"Columns relevant to your prompt: {', '.join(matched_cols)}")

    # 2) Value hotspots per matched column
    for c in matched_cols[:6]:  # cap to avoid noise
        try:
            vc = df[c].astype(str).value_counts().head(5)
            top_vals = [f"{idx} ({int(cnt)})" for idx, cnt in vc.items()]
            highlights.append(f"Top values in **{c}**: " + ", ".join(top_vals))
        except Exception:
            pass

    # 3) Simple row filters for binary/flag-like columns containing keywords
    for c in matched_cols:
        s = df[c]
        # Look for common bad/alert tokens
        alert_tokens = ["delayed", "late", "overdue", "error", "failed", "critical", "pending", "open"]
        mask = s.astype(str).str.lower().str.contains("|".join(alert_tokens), na=False)
        count_alert = int(mask.sum())
        if count_alert:
            highlights.append(f"Found {count_alert} rows in **{c}** with potential issues (e.g., {', '.join(alert_tokens[:4])}).")

    return highlights


def generate_review_points(df: pd.DataFrame, profile: dict, prompt_text: str):
    """Create a clean list of review points from profile + prompt guidance."""
    points = []

    # Data overview
    points.append(f"Loaded {profile['rows']:,} rows, {profile['cols']} columns from {len(profile.get('files', []))} file(s).")

    # Missing data
    high_missing = [c for c in profile["columns"] if c["null_pct"] >= 10]
    if high_missing:
        names = ", ".join([f"{c['name']} ({c['null_pct']}%)" for c in sorted(high_missing, key=lambda x: -x['null_pct'])[:10]])
        points.append(f"Columns with ≥10% missing values: {names}.")

    # Duplicates
    if profile["duplicates"] > 0:
        points.append(f"Detected {profile['duplicates']:,} duplicate row(s). Consider deduplication using potential keys: {', '.join(profile['potential_keys']) or 'none found'}.")

    # Outliers in numeric columns
    outlier_cols = [(c["name"], c["outliers"]) for c in profile["columns"] if c["outliers"] > 0]
    if outlier_cols:
        top = ", ".join([f"{n}: {cnt}" for n, cnt in sorted(outlier_cols, key=lambda x: -x[1])[:10]])
        points.append(f"Outliers detected (IQR method) → {top}.")

    # Date anomalies
    today = date.today().isoformat()
    date_cols = [c for c in profile["columns"] if c["is_date"]]
    for c in date_cols:
        if c["date_min"] and c["date_max"]:
            # Flag future dates
            try:
                if c["date_max"] > today:
                    points.append(f"Date column **{c['name']}** has future dates up to {c['date_max']}.")
            except Exception:
                pass

    # Column cardinality suggestions
    high_card = [c for c in profile["columns"] if c["unique"] > max(1, int(0.9 * profile["rows"]))]
    if high_card:
        points.append("High-cardinality columns (likely identifiers) → " + ", ".join([c["name"] for c in high_card[:10]]) + ".")

    # Prompt-aligned highlights
    keywords = extract_keywords(prompt_text)
    points.extend(prompt_aligned_highlights(df, keywords))

    # Keep list concise
    deduped = []
    seen = set()
    for p in points:
        if p not in seen:
            deduped.append(p)
            seen.add(p)

    return deduped


def format_profile_for_display(profile: dict):
    lines = []
    lines.append(f"Rows: {profile['rows']:,}")
    lines.append(f"Columns: {profile['cols']}")
    lines.append(f"Duplicate rows: {profile['duplicates']:,}")
    if profile.get("files"):
        lines.append("Files: " + ", ".join(profile["files"]))
    return "
".join(lines)


def build_docx(review_points, profile):
    doc = Document()
    doc.add_heading('Excel Review - Summary', 0)

    doc.add_heading('Overview', level=1)
    doc.add_paragraph(format_profile_for_display(profile))

    doc.add_heading('Review Points', level=1)
    for i, p in enumerate(review_points, start=1):
        doc.add_paragraph(f"{i}. {p}")

    # Column-level appendix
    doc.add_heading('Column Appendix', level=1)
    for col in profile["columns"]:
        para = doc.add_paragraph()
        para.add_run(col["name"]).bold = True
        details = []
        details.append(f"dtype={col['dtype']}")
        details.append(f"nulls={col['nulls']} ({col['null_pct']}%)")
        details.append(f"unique={col['unique']}")
        if col["numeric_stats"]:
            ns = col["numeric_stats"]
            details.append(f"min={ns['min']}")
            details.append(f"max={ns['max']}")
        if col["is_date"]:
            details.append(f"date_range={col['date_min']} → {col['date_max']}")
        details.append(f"outliers={col['outliers']}")
        doc.add_paragraph("; ".join(details))
        if col["top_values"]:
            doc.add_paragraph("Top values: " + ", ".join(col["top_values"]))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ---------------------------
# Streamlit App
# ---------------------------
st.set_page_config(page_title="Excel Review Dashboard", layout="wide")

st.title("📊 Excel Review Dashboard")

st.markdown(
    """
    Upload up to **5** Excel/CSV files, enter a prompt (what you want to review), and click **Analyze**.
    The app will merge the data (outer columns), run quick data quality checks, and generate **review points**
    aligned with your prompt.

    **Tip:** Prompts can be things like _"find delays by region", "flag overdue POs", "highlight anomalies in cost"_.
    """
)

uploaded = st.file_uploader(
    "Upload Excel/CSV files (max 5)",
    type=["xlsx", "xls", "xlsm", "csv"],
    accept_multiple_files=True,
    help="First visible sheet is read for Excel files."
)

prompt_text = st.text_area("Your prompt (what should we review?)", height=80)

col_run, col_opt = st.columns([1,1])
with col_run:
    run = st.button("Analyze", type="primary")

with col_opt:
    enable_llm = st.checkbox("Use LLM (OpenAI) for refined review points (optional)")
    api_key = None
    if enable_llm:
        st.info("Optional: Provide your **OpenAI API key** (will be used only during this session; not stored).")
        api_key = st.text_input("OpenAI API Key", type="password")

if run:
    if not uploaded:
        st.warning("Please upload at least one file.")
        st.stop()

    files = uploaded[:5]
    frames = []
    for uf in files:
        df = read_tabular(uf)
        if df is not None:
            frames.append(df)

    if not frames:
        st.error("No readable files were provided.")
        st.stop()

    # Outer-join by columns (stack rows, align disparate schemas)
    try:
        df_all = pd.concat(frames, ignore_index=True, sort=False)
    except Exception as e:
        st.error(f"Failed to concatenate files: {e}")
        st.stop()

    # Profile and review points
    profile = profile_dataframe(df_all)
    review_points = generate_review_points(df_all, profile, prompt_text)

    # Optional: augment using LLM if key provided
    if enable_llm and api_key and len(review_points) > 0:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=api_key)
            # Keep the prompt small to avoid token issues; send the top profile lines + current points
            sys_msg = (
                "You are a data QA and business review assistant. Based on the provided summary and user prompt, "
                "rewrite and refine the review points to be crisp, prioritized, and actionable. If evidence is weak, "
                "add a short note on what extra data would help. Output as a bullet list (5-12 bullets)."
            )
            summary = format_profile_for_display(profile)
            bullets = "
".join([f"- {p}" for p in review_points])
            user_msg = f"Prompt: {prompt_text}

Data summary:
{summary}

Current review points:
{bullets}"
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": sys_msg},
                    {"role": "user", "content": user_msg},
                ],
                temperature=0.2,
            )
            refined = resp.choices[0].message.content.strip()
            st.success("LLM-refined review points:")
            st.markdown(refined)
        except Exception as e:
            st.warning(f"LLM refinement skipped due to error: {e}")

    # UI: Tabs
    tab1, tab2, tab3 = st.tabs(["Preview", "Data Quality", "Review Points"])

    with tab1:
        st.subheader("Merged Preview")
        st.dataframe(df_all.head(50))
        st.caption(f"Showing first 50 of {len(df_all):,} rows. Column 'source_file' indicates origin.")

    with tab2:
        st.subheader("Quick Profile")
        st.text(format_profile_for_display(profile))

        st.divider()
        st.subheader("Column Details")
        # Show a compact table of columns
        cols_df = pd.DataFrame(profile["columns"])[:, ["name", "dtype", "nulls", "null_pct", "unique", "outliers", "is_date", "date_min", "date_max"]] if False else pd.DataFrame(profile["columns"]).loc[:, ["name", "dtype", "nulls", "null_pct", "unique", "outliers", "is_date", "date_min", "date_max"]]
        st.dataframe(cols_df)

    with tab3:
        st.subheader("Review Points")
        if review_points:
            for i, p in enumerate(review_points, start=1):
                st.markdown(f"**{i}.** {p}")
        else:
            st.info("No specific review points generated yet. Try adding a more detailed prompt.")

        # Export DOCX
        if st.button("Export to Word (.docx)"):
            bio = build_docx(review_points, profile)
            st.download_button(
                label="Download review_summary.docx",
                data=bio,
                file_name="review_summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
